from fpdf import FPDF
from docx import Document
import streamlit as st
from io import BytesIO

# Function to save cover letter as PDF
def save_as_pdf(content):
    # Initialize PDF object
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Split content into lines for writing
    lines = content.split("\n")
    for line in lines:
        pdf.cell(0, 10, line, ln=True)

    # Output PDF to a string and write to BytesIO buffer
    pdf_buffer = BytesIO()
    pdf_data = pdf.output(dest="S").encode("latin1")  # Output as string and encode to bytes
    pdf_buffer.write(pdf_data)
    pdf_buffer.seek(0)

    return pdf_buffer.getvalue()  # Return the data as bytes

# Function to save cover letter as Word document
def save_as_docx(content):
    doc = Document()
    doc.add_heading("Cover Letter", level=1)

    # Add the cover letter content
    for line in content.split("\n"):
        doc.add_paragraph(line)

    # Save to a BytesIO object
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer.getvalue()  # Return the data as bytes